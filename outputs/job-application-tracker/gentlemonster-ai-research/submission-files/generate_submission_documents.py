from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESUME_SOURCE = ROOT / "gentlemonster-ai-experience-resume-content-final-2026-07-12.md"
COVER_SOURCE = ROOT / "gentlemonster-ai-experience-cover-letter-final-2026-07-12.md"

RESUME_DOCX = ROOT / "Yuminseok_Yu_GENTLE_MONSTER_AI_Experience_Resume_2026-07-12.docx"
COVER_DOCX = ROOT / "Yuminseok_Yu_GENTLE_MONSTER_AI_Experience_Cover_Letter_2026-07-12.docx"

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
ACCENT = RGBColor(91, 111, 34)
LIGHT = "D9DEE6"
FONT_NAME = "AppleGothic"


def clean_inline(value: str) -> str:
    value = value.strip()
    value = value.replace("`", "").replace("**", "")
    return re.sub(r"\s{2,}", " ", value)


def set_font(run, *, size: float, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def paragraph_rule(paragraph, color: str = LIGHT, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    set_font(run, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(document: Document, *, compact: bool) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.1 if compact else 2.05)
    section.bottom_margin = Cm(1.0 if compact else 2.05)
    section.left_margin = Cm(1.25 if compact else 2.15)
    section.right_margin = Cm(1.25 if compact else 2.15)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(8.4 if compact else 10.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2.5 if compact else 9)
    normal.paragraph_format.line_spacing = 1.1 if compact else 1.52

    for name, size in (("Title", 25), ("Heading 1", 12.5), ("Heading 2", 10.6)):
        style = document.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = INK
        style.font.bold = True

    footer = section.footer.paragraphs[0]
    add_page_field(footer)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "5B6F22")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_resume_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title.upper())
    set_font(run, size=9.7, bold=True, color=ACCENT)
    paragraph_rule(paragraph, color="9AA676", size=10)


def add_project_heading(document: Document, title: str, meta: str | None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(clean_inline(title))
    set_font(run, size=9.4, bold=True)
    if meta:
        meta_paragraph = document.add_paragraph()
        meta_paragraph.paragraph_format.space_after = Pt(1)
        meta_paragraph.paragraph_format.keep_with_next = True
        meta_run = meta_paragraph.add_run(clean_inline(meta))
        set_font(meta_run, size=7.5, bold=True, color=MUTED)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.48)
    paragraph.paragraph_format.first_line_indent = Cm(-0.2)
    paragraph.paragraph_format.space_after = Pt(0.8)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(clean_inline(text))
    set_font(run, size=8.0, color=INK)


def build_resume() -> None:
    lines = RESUME_SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document, compact=True)
    document.core_properties.title = "GENTLE MONSTER AI Experience Resume - Yuminseok Yu"
    document.core_properties.author = "Yuminseok Yu"

    title = clean_inline(lines[0].removeprefix("# "))
    role = clean_inline(lines[2])
    contact = clean_inline(lines[3])

    header = document.add_paragraph()
    header.paragraph_format.space_after = Pt(2)
    run = header.add_run(title)
    set_font(run, size=22, bold=True)

    role_paragraph = document.add_paragraph()
    role_paragraph.paragraph_format.space_after = Pt(3)
    role_run = role_paragraph.add_run(role)
    set_font(role_run, size=10.5, bold=True, color=ACCENT)

    contact_paragraph = document.add_paragraph()
    contact_paragraph.paragraph_format.space_after = Pt(4)
    contact_run = contact_paragraph.add_run(contact)
    set_font(contact_run, size=8.0, color=MUTED)
    paragraph_rule(contact_paragraph)

    index = 4
    pending_subheading: str | None = None
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("## "):
            add_resume_section_heading(document, line[3:])
            pending_subheading = None
        elif line.startswith("### "):
            pending_subheading = line[4:]
            meta = None
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines):
                candidate = lines[lookahead].strip()
                if candidate and not candidate.startswith(("#", "-")):
                    meta = candidate
                    index = lookahead
            add_project_heading(document, pending_subheading, meta)
            pending_subheading = None
        elif line.startswith("- "):
            link_match = re.match(r"- ([^:]+): (https?://\S+)$", clean_inline(line[2:]))
            if link_match:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1)
                label_run = paragraph.add_run(f"{link_match.group(1)}: ")
                set_font(label_run, size=8.0, bold=True)
                add_hyperlink(paragraph, link_match.group(2), link_match.group(2))
            else:
                add_bullet(document, line[2:])
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(clean_inline(line))
            set_font(run, size=8.2, color=INK)
        index += 1

    document.save(RESUME_DOCX)


def build_cover_letter() -> None:
    lines = COVER_SOURCE.read_text(encoding="utf-8").splitlines()
    title = clean_inline(lines[0].removeprefix("# "))
    paragraphs = [clean_inline(line) for line in lines[1:] if line.strip()]

    document = Document()
    configure_document(document, compact=False)
    section = document.sections[0]
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    document.core_properties.title = "GENTLE MONSTER AI Experience Cover Letter - Yuminseok Yu"
    document.core_properties.author = "Yuminseok Yu"

    top = document.add_paragraph()
    top.paragraph_format.space_after = Pt(3)
    top_run = top.add_run("GENTLE MONSTER  |  AI EXPERIENCE")
    set_font(top_run, size=9.2, bold=True, color=ACCENT)
    paragraph_rule(top, color="9AA676", size=12)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(3)
    heading_run = heading.add_run(title)
    set_font(heading_run, size=21, bold=True)

    applicant = document.add_paragraph()
    applicant.paragraph_format.space_after = Pt(12)
    applicant_run = applicant.add_run("유민석  |  AI Experience / AI Creative")
    set_font(applicant_run, size=9.5, bold=True, color=MUTED)

    for body in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.35
        run = paragraph.add_run(body)
        set_font(run, size=9.6)

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(4)
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = closing.add_run("유민석")
    set_font(run, size=9.6, bold=True, color=ACCENT)

    document.save(COVER_DOCX)


if __name__ == "__main__":
    build_resume()
    build_cover_letter()
    print(RESUME_DOCX)
    print(COVER_DOCX)
