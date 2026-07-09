from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUT_DIR = Path("/Users/yuminseog/portfolio/outputs/job-application-tracker/w3company")
OUT_DIR.mkdir(parents=True, exist_ok=True)

CONTACT_EMAIL = "aheayabaraya@gmail.com"
PHONE = "+82 10-4229-9239"

LINKS = {
    "Portfolio": "https://portfolio-yuminseoks-projects.vercel.app",
    "Loom/Pulso": "https://loom-signal-deck.vercel.app/",
    "AHEYA": "https://aheyabaraya.xyz/",
}

cv = {
    "name": "Minseok Yu",
    "role": "AI Content Marketer Intern | AI Creative / Content Marketing",
    "contact": f"Seoul, Korea | {PHONE} | {CONTACT_EMAIL}",
    "profile": [
        "Entry-level AI creative and content marketing candidate with a Fashion Marketing background and portfolio projects in AI short-form ads, AI idol/IP content, and product-facing GTM content.",
        "Strong fit for W3Company's AI Content Marketer role: I understand Moji as a casual global friendship/language-exchange app and SeriUs as a trust-based serious international dating app, then translate each service tone into social/ad creative directions.",
        "My working style is AI-native planning first: concept, storyboard, prompt system, asset generation, review, revision, and social/ad-ready output. I build and compare creative candidates to learn what hooks, scenes, and formats can drive engagement."
    ],
    "skills": [
        "AI-native content planning: concept boards, campaign hooks, storyboards, short-form ad flow, copy direction, service-tone adaptation.",
        "Generative AI workflow: Grok/GPT/Codex workflow, Suno, Seedance 2.0, Kling, Grok Imagine, Midjourney, Gemini, Nano Banana.",
        "Production and publishing tools: Adobe Premiere Pro, Photoshop, CapCut, web/portfolio publishing, YouTube/TikTok/X content routing.",
        "Marketing perspective: campaign message structure, app/service public content, CTA flow, KPI/event tracking design, feedback-based iteration.",
        "English communication: SFTI-CMU English abstract and poster-style research communication."
    ],
    "projects": [
        {
            "title": "MUSINSA Mujinjang AI Short-form Ad",
            "meta": "Competition project, 3-person team | Jun 2026",
            "bullets": [
                "Translated the campaign message \"Break bias, wear diversity, meet Mujinjang\" into a 30-second AI short-form ad concept.",
                "Built message structure, main storyboard, scene order, and production flow so the ad stayed campaign-led rather than just visually generated.",
                "Generated and reviewed image/video candidates with GPT, Seedance 2.0, Kling, and Grok, then adjusted cut flow and transitions for a social-first rhythm."
            ]
        },
        {
            "title": "Loom / Pulso AI Idol IP Content System",
            "meta": "Personal AI creative project | May 2026 - Jun 2026",
            "bullets": [
                "Designed a 13-member AI idol IP system across member identities, tracks, MV-style scenes, short-form clips, web archive, and fan-participation surfaces.",
                "Built a repeatable AI production loop using Grok/GPT/Codex for planning and prompt structuring, batch image candidate generation, Grok Imagine video candidates, and CapCut/Codex edit checks.",
                "Expanded Pulso from a core video concept into TikTok/X-style short-form cuts and track-board content for audience-facing distribution."
            ]
        },
        {
            "title": "ADSB AI-assisted Fashion Brand Short-form",
            "meta": "Academic-industry collaboration | Sep 2025 - Nov 2025",
            "bullets": [
                "Researched Andersson Bell's brand mood and converted it into shot flow, storyboard direction, and a 15-second AI-assisted short-form proof.",
                "Created and revised image/video clip candidates using Midjourney, Gemini, Nano Banana, Photoshop, Kling, and Adobe Premiere Pro after practitioner feedback.",
                "Focused on matching brand tone and visual rhythm, not only on tool output quality."
            ]
        },
        {
            "title": "AHEYA Product GTM / Public Content",
            "meta": "Personal service prototype | Sep 2025 - Mar 2026",
            "bullets": [
                "Structured public-page messaging, X post concepts, CTA flow, and character/world visual assets for a Web3 service prototype.",
                "Converted complex wallet/testnet participation into user-facing actions such as support, feedback, share, and save.",
                "Designed KPI/event tracking logic around visit, support intent, feedback, share, and saved-record events so content and product decisions could be reviewed without overstating growth results."
            ]
        },
        {
            "title": "SFTI-CMU Research Communication",
            "meta": "International conference material | May 2025 - Jun 2025",
            "bullets": [
                "Developed the English abstract and poster-style structure for \"AI-Generated Emotional Content Strategy for Niche Fashion Brands: Toward Identity-Based Visual Clustering.\"",
                "This framing later informed AHEYA's visual identity and emotional content direction."
            ]
        }
    ],
    "experience": [
        {
            "title": "Teaching Assistant, Global Fashion Industry Department",
            "meta": "Hansung University | Sep 2024 - Dec 2025",
            "bullets": [
                "Supported class operations, studio/classroom materials, student inquiries, documents, and academic administration.",
                "Helped new faculty classes run smoothly by organizing communication and class preparation details."
            ]
        }
    ],
    "education": [
        "Hansung University, B.A. Fashion Marketing, expected Aug 2026"
    ],
    "languages": [
        "Korean: Native",
        "English: Conversational, OPic IM, Apr 23 2025"
    ]
}

cover_sections = {
    "지원동기 / 회사 관심 이유": (
        "더블유쓰리컴퍼니는 Moji와 SeriUs를 통해 국경과 언어를 넘는 연결 경험을 만들고 있습니다. "
        "두 서비스는 모두 매칭과 대화를 다루지만, Moji는 글로벌 친구·언어교환·가벼운 vibe 매칭에 "
        "가깝고, SeriUs는 검증과 신뢰를 바탕으로 한 진지한 국제 관계 서비스에 가깝다고 이해했습니다. "
        "따라서 AI 콘텐츠 마케터 인턴 역할은 단순히 생성형 AI 결과물을 많이 만드는 일이 아니라, "
        "서비스별 톤을 구분하고 사용자가 멈춰 볼 만한 숏폼, 이미지, 오디오, 광고 소재로 빠르게 "
        "실험하는 일이라고 봤습니다. 저는 패션마케팅 전공과 AI 숏폼/브랜드 콘텐츠 프로젝트를 통해 "
        "메시지 구조, 장면 흐름, 시각 결과물, 공개 동선을 함께 설계해왔고, 이 경험을 W3의 글로벌 "
        "앱 콘텐츠 실험에 연결하고 싶습니다."
    ),
    "직무 적합 경험": (
        "정규 경력은 신입이지만, MUSINSA, Loom/Pulso, ADSB, AHEYA를 통해 AI를 활용한 콘텐츠 "
        "기획과 제작 루프를 직접 실험해왔습니다. 무신사 AI 광고제에서는 '편견을 벗다, 다양성을 "
        "입다, 무진장을 만나다'라는 메시지를 30초 AI 숏폼 광고로 만들기 위해 스토리보드, 장면 "
        "순서, 컷 흐름을 먼저 구조화했고, Seedance 2.0, Kling, Grok 기반 이미지·영상 후보를 "
        "검토하며 광고 메시지가 흐려지지 않도록 리듬을 조정했습니다. Loom/Pulso에서는 Grok, "
        "GPT, Codex를 연결해 기획과 프롬프트를 구조화하고, 이미지·영상 후보 생성, Grok Imagine "
        "영상화, CapCut/Codex 편집 확인으로 이어지는 반복 가능한 제작 루프를 만들었습니다. ADSB "
        "산학 프로젝트에서는 Photoshop과 Adobe Premiere Pro를 활용해 AI 이미지·영상 후보를 "
        "15초 숏폼 영상의 컷 흐름으로 연결했습니다. AHEYA에서는 방문, 후원 의도, 피드백, 공유, "
        "저장처럼 어떤 이벤트를 봐야 다음 콘텐츠/제품 판단으로 이어질 수 있는지 KPI/event tracking "
        "구조를 설계했고, SFTI-CMU에서는 AI-generated emotional content strategy를 영문 초록과 "
        "포스터형 구조로 정리했습니다."
    ),
    "입사 후 기여 / 실행 계획": (
        "입사 후에는 글로벌 Gen-Z 트렌드를 이미 완성형으로 안다고 단정하기보다, 빠르게 후보를 "
        "만들고 비교하며 반응 가능한 포맷을 찾아가는 방식으로 기여하겠습니다. Moji에는 언어교환, "
        "외국인 친구, 스와이프 매칭, 대화 시작 상황을 가볍고 빠른 숏폼 훅으로 풀고, SeriUs에는 "
        "검증, 신뢰, 번역, 진지한 국제 관계라는 메시지를 더 차분하고 신뢰감 있는 광고 소재로 "
        "나누어 실험하겠습니다. 각 콘셉트는 이미지, 숏폼 영상, 오디오, 카피 후보로 쪼개고, Kling, "
        "Seedance 2.0, Grok 계열 워크플로우와 새로운 AI 툴을 빠르게 붙여보겠습니다. 단순히 많이 "
        "만드는 데서 끝내지 않고, 첫 장면, 문장, 시각 톤, CTA, 저장·공유·클릭 같은 판단 기준을 "
        "함께 정리해 마케팅팀이 다음 캠페인 소재를 고를 수 있는 실험 기록으로 남기겠습니다."
    )
}

integrated_cover = "\n\n".join(cover_sections.values())

email_subject = "[AI Content Marketer Intern 지원] 유민석"
email_body = f"""안녕하세요, 더블유쓰리컴퍼니 AI 콘텐츠 마케터 인턴 포지션에 지원하는 유민석입니다.

저는 AI로 숏폼 광고, AI 아이돌/IP 콘텐츠, 제품 공개/마케팅 콘텐츠를 기획하고 실제 산출물로 제작해왔습니다. Moji는 글로벌 친구·언어교환 중심의 가벼운 매칭 서비스, SeriUs는 검증과 신뢰 기반의 serious international dating 서비스로 이해했고, 두 서비스의 톤을 나누어 AI 숏폼·이미지·오디오 광고 소재를 빠르게 실험하는 역할에 강하게 맞는다고 판단해 지원드립니다.

영문 CV와 포트폴리오 링크를 함께 전달드립니다.
- Portfolio: {LINKS["Portfolio"]}
- Loom/Pulso: {LINKS["Loom/Pulso"]}
- AHEYA: {LINKS["AHEYA"]}

감사합니다.
유민석 드림
{PHONE}
{CONTACT_EMAIL}
"""

validation_note = """# Wanted Resume Check for W3Company

Verdict: broadly aligned, but the current Wanted PDF is weaker than the latest portfolio/master positioning for W3Company.

What matches:
- Main portfolio pillars are consistent: MUSINSA, Loom/Pulso, ADSB, SFTI-CMU, AHEYA.
- The resume correctly positions you as entry-level/new graduate, not as someone with formal industry employment.
- The core direction is right: AI is used for planning, storyboarding, image/video generation, and content output.

What should be strengthened for W3:
- Put global Gen-Z social formats, memes, culturally resonant concepts, and app campaign creatives closer to the top.
- Make the AI production loop clearer: LLM planning -> prompt registry -> image/video candidates -> review/revision -> short-form/ad-ready output.
- Reframe editing as a supporting skill. W3 should see you as AI creative/content marketing first, not a pure editor.
- Add readiness to experiment with tools named in the posting such as Veo3 and ElevenLabs, while only claiming direct experience with tools actually used.

Submission note:
- Email spelling is copied from the Wanted PDF: aheayabaraya@gmail.com. Because it differs from the AHEYA project spelling, verify before sending.
"""


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_docx_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color=(46, 116, 181))
    return p


def add_docx_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_docx_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_font(run)
    return p


def build_docx(path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(cv["name"])
    set_font(run, size=22, bold=True, color=(11, 37, 69))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run(cv["role"])
    set_font(run, size=12, bold=True, color=(31, 77, 120))

    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(8)
    run = contact.add_run(cv["contact"])
    set_font(run, size=10, color=(85, 85, 85))

    add_docx_heading(doc, "Profile")
    for item in cv["profile"]:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Core Skills")
    for item in cv["skills"]:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Selected AI Creative and Marketing Projects")
    for project in cv["projects"]:
        add_docx_body(doc, project["title"], bold_prefix=project["title"])
        add_docx_body(doc, project["meta"])
        for item in project["bullets"]:
            add_docx_bullet(doc, item)

    add_docx_heading(doc, "Experience")
    for exp in cv["experience"]:
        add_docx_body(doc, exp["title"], bold_prefix=exp["title"])
        add_docx_body(doc, exp["meta"])
        for item in exp["bullets"]:
            add_docx_bullet(doc, item)

    add_docx_heading(doc, "Education")
    for item in cv["education"]:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Languages")
    for item in cv["languages"]:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Links")
    for label, url in LINKS.items():
        add_docx_body(doc, f"{label}: {url}")

    doc.core_properties.title = "Minseok Yu - AI Content Marketer Intern CV"
    doc.core_properties.author = "Minseok Yu"
    doc.save(path)


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "Name",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=3,
            alignment=TA_LEFT,
        ),
        "role": ParagraphStyle(
            "Role",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=3,
        ),
        "meta": ParagraphStyle(
            "Meta",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#555555"),
            spaceAfter=8,
        ),
        "h": ParagraphStyle(
            "Heading",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=14,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
            spaceAfter=3,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
            leftIndent=14,
            firstLineIndent=-7,
            bulletIndent=0,
            spaceAfter=2.5,
        ),
    }


def add_pdf_bullet(story, text, styles):
    story.append(Paragraph(text, styles["bullet"], bulletText="-"))


def build_pdf(path: Path):
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title="Minseok Yu - AI Content Marketer Intern CV",
        author="Minseok Yu",
    )
    story = [
        Paragraph(cv["name"], styles["name"]),
        Paragraph(cv["role"], styles["role"]),
        Paragraph(cv["contact"], styles["meta"]),
    ]

    for heading, items in [
        ("Profile", cv["profile"]),
        ("Core Skills", cv["skills"]),
    ]:
        story.append(Paragraph(heading, styles["h"]))
        for item in items:
            add_pdf_bullet(story, item, styles)

    story.append(Paragraph("Selected AI Creative and Marketing Projects", styles["h"]))
    for project in cv["projects"]:
        story.append(Paragraph(f"<b>{project['title']}</b>", styles["body"]))
        story.append(Paragraph(project["meta"], styles["meta"]))
        for item in project["bullets"]:
            add_pdf_bullet(story, item, styles)
        story.append(Spacer(1, 2))

    story.append(Paragraph("Experience", styles["h"]))
    for exp in cv["experience"]:
        story.append(Paragraph(f"<b>{exp['title']}</b>", styles["body"]))
        story.append(Paragraph(exp["meta"], styles["meta"]))
        for item in exp["bullets"]:
            add_pdf_bullet(story, item, styles)

    story.append(Paragraph("Education", styles["h"]))
    for item in cv["education"]:
        add_pdf_bullet(story, item, styles)

    story.append(Paragraph("Languages", styles["h"]))
    for item in cv["languages"]:
        add_pdf_bullet(story, item, styles)

    story.append(Paragraph("Links", styles["h"]))
    for label, url in LINKS.items():
        story.append(Paragraph(f"<b>{label}:</b> {url}", styles["body"]))

    doc.build(story)


def build_markdown(path: Path):
    lines = [
        f"# {cv['name']}",
        "",
        f"**{cv['role']}**",
        "",
        cv["contact"],
        "",
        "## Profile",
    ]
    for item in cv["profile"]:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## Core Skills")
    for item in cv["skills"]:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## Selected AI Creative and Marketing Projects")
    for project in cv["projects"]:
        lines.extend(["", f"### {project['title']}", project["meta"]])
        for item in project["bullets"]:
            lines.append(f"- {item}")
    lines.extend(["", "## Experience"])
    for exp in cv["experience"]:
        lines.extend(["", f"### {exp['title']}", exp["meta"]])
        for item in exp["bullets"]:
            lines.append(f"- {item}")
    lines.extend(["", "## Education"])
    for item in cv["education"]:
        lines.append(f"- {item}")
    lines.extend(["", "## Languages"])
    for item in cv["languages"]:
        lines.append(f"- {item}")
    lines.extend(["", "## Links"])
    for label, url in LINKS.items():
        lines.append(f"- {label}: {url}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_text_files():
    cover = OUT_DIR / "w3company-cover-letter-ko.md"
    cover_lines = [
        "# 더블유쓰리컴퍼니 AI 콘텐츠 마케터 인턴 자기소개서",
        "",
        "## 제출용 통합본",
        "",
        integrated_cover,
        "",
        "## 섹션별 원문",
        "",
    ]
    for title, text in cover_sections.items():
        cover_lines.extend([f"### {title}", "", text, ""])
    cover.write_text("\n".join(cover_lines), encoding="utf-8")

    email = OUT_DIR / "w3company-email-draft.md"
    email.write_text(
        f"# W3Company Email Draft\n\nTo: weare@w3company.team\n\nSubject: {email_subject}\n\n{email_body}",
        encoding="utf-8",
    )

    validation = OUT_DIR / "wanted-resume-validation-for-w3.md"
    validation.write_text(validation_note, encoding="utf-8")

    return cover, email, validation


def main():
    cv_md = OUT_DIR / "Minseok_Yu_AI_Content_Marketer_CV.md"
    cv_docx = OUT_DIR / "Minseok_Yu_AI_Content_Marketer_CV.docx"
    cv_pdf = OUT_DIR / "Minseok_Yu_AI_Content_Marketer_CV.pdf"

    build_markdown(cv_md)
    build_docx(cv_docx)
    build_pdf(cv_pdf)
    cover, email, validation = write_text_files()

    manifest = {
        "company": "더블유쓰리컴퍼니",
        "role": "AI 콘텐츠 마케터 (인턴)",
        "created_at": "2026-06-19",
        "files": {
            "cv_markdown": str(cv_md),
            "cv_docx": str(cv_docx),
            "cv_pdf": str(cv_pdf),
            "cover_letter_ko": str(cover),
            "email_draft": str(email),
            "validation": str(validation),
        },
        "cover_sections": cover_sections,
        "integrated_cover": integrated_cover,
        "email_subject": email_subject,
        "email_body": email_body,
    }
    manifest_path = OUT_DIR / "w3company-application-manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
